from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Dict, List, Tuple

from docx import Document


@dataclass
class ParsedDocument:
    """Result of parsing a simple Markdown/text document."""

    title: str
    fields: Dict[str, str]
    warnings: List[str]
    raw_text: str = ""


def normalize_field_name(name: str) -> str:
    return name.strip().lower().replace(" ", "_")


def normalize_text_line_breaks(text: str) -> str:
    """Normalize line-break representations from Office/text sources.

    - CRLF/CR -> LF
    - Literal `_x000D_` marker -> LF
    """

    return text.replace("\r\n", "\n").replace("\r", "\n").replace("_x000D_", "\n")


def parse_kv_markdown(markdown_text: str) -> ParsedDocument:
    """Parse lightweight key-value markdown for fast business input.

    Supported format examples:
    - title: Q2 销售复盘
    - summary: 本季度营收同比增长 18%

    Lines without `:` are ignored and returned as warnings.
    """

    markdown_text = normalize_text_line_breaks(markdown_text)

    title = ""
    fields: Dict[str, str] = {}
    warnings: List[str] = []
    current_key = ""

    for line_no, raw_line in enumerate(markdown_text.splitlines(), start=1):
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# ") and not title:
            title = stripped[2:].strip()
            continue

        sep_positions = [pos for pos in (line.find(":"), line.find("：")) if pos >= 0]
        if not sep_positions:
            if current_key:
                # Continuation lines are treated as part of the previous field value.
                fields[current_key] = (fields[current_key] + "\n" + stripped).strip()
            else:
                warnings.append(f"第 {line_no} 行未包含冒号，已忽略: {stripped}")
            continue

        split_idx = min(sep_positions)
        key = line[:split_idx]
        value = line[split_idx + 1 :]
        normalized_key = normalize_field_name(key)
        if not normalized_key:
            warnings.append(f"第 {line_no} 行的字段名为空，已忽略")
            current_key = ""
            continue
        fields[normalized_key] = value.strip()
        current_key = normalized_key

    return ParsedDocument(title=title, fields=fields, warnings=warnings, raw_text=markdown_text)


def extract_text_from_uploaded_file(file_name: str, file_bytes: bytes) -> Tuple[str, List[str]]:
    """Extract plain text from uploaded files.

    Supported extensions: .txt, .md, .docx
    """

    warnings: List[str] = []
    suffix = file_name.lower().rsplit(".", 1)[-1] if "." in file_name else ""

    if suffix in {"txt", "md"}:
        for encoding in ("utf-8", "utf-8-sig", "gbk"):
            try:
                return normalize_text_line_breaks(file_bytes.decode(encoding)), warnings
            except UnicodeDecodeError:
                continue
        warnings.append("文本文件编码无法识别，请使用 UTF-8 或 GBK 编码。")
        return "", warnings

    if suffix == "docx":
        try:
            doc = Document(io.BytesIO(file_bytes))
            lines = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            return normalize_text_line_breaks("\n".join(lines)), warnings
        except Exception:
            warnings.append("DOCX 解析失败，请检查文件是否损坏。")
            return "", warnings

    warnings.append("仅支持 .txt / .md / .docx 文件。")
    return "", warnings


def infer_fields_from_plain_text(raw_text: str, target_tokens: List[str]) -> ParsedDocument:
    """Infer field mapping from plain paragraphs without `key: value` format.

    Strategy:
    - First line -> title-like token if available, else first token.
    - For each line, try keyword match with token name.
    - Unmatched lines fill remaining tokens in order.
    """

    raw_text = normalize_text_line_breaks(raw_text)
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    fields: Dict[str, str] = {}
    warnings: List[str] = []

    if not lines:
        return ParsedDocument(title="", fields=fields, warnings=warnings, raw_text=raw_text)

    if not target_tokens:
        warnings.append("模板中未扫描到文本占位符，无法自动映射文本内容。")
        return ParsedDocument(title=lines[0], fields=fields, warnings=warnings, raw_text=raw_text)

    normalized_tokens = [normalize_field_name(token) for token in target_tokens]
    available_tokens = list(normalized_tokens)

    # Prefer common title tokens for the first paragraph.
    title_candidates = [
        token
        for token in available_tokens
        if any(k in token for k in ["title", "project_name", "name", "主题", "标题"])
    ]
    if title_candidates:
        fields[title_candidates[0]] = lines[0]
        available_tokens.remove(title_candidates[0])
        start_idx = 1
    else:
        fields[available_tokens[0]] = lines[0]
        available_tokens.pop(0)
        start_idx = 1

    for line in lines[start_idx:]:
        matched = None
        for token in list(available_tokens):
            token_parts = [part for part in token.split("_") if len(part) >= 2]
            if any(part in line.lower() for part in token_parts):
                matched = token
                break

        if matched:
            fields[matched] = line
            available_tokens.remove(matched)
            continue

        if available_tokens:
            token = available_tokens.pop(0)
            fields[token] = line
        else:
            warnings.append(f"存在未映射段落，已忽略: {line[:30]}")

    if available_tokens:
        warnings.append("部分模板字段未能从自然段文案中自动匹配，请手动补充。")

    return ParsedDocument(title=lines[0], fields=fields, warnings=warnings, raw_text=raw_text)
